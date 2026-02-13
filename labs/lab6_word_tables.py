#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Лабораторная работа №6
Microsoft Office Word. Создание и форматирование таблиц
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn

def create_lab6_document():
    """Создание документа с таблицами для лабораторной работы №6"""
    
    # Создание нового документа
    doc = Document()
    
    # Установка параметров страницы
    section = doc.sections[0]
    section.page_height = Cm(29.7)  # A4
    section.page_width = Cm(21)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(1)
    
    # Заголовок
    title = doc.add_paragraph()
    title_run = title.add_run("Лабораторная работа №6\nСоздание и форматирование таблиц")
    title_run.font.bold = True
    title_run.font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Название таблицы
    table_title = doc.add_paragraph()
    table_title_run = table_title.add_run("Таблица 1. Автобусные маршруты")
    table_title_run.font.bold = True
    table_title_run.font.size = Pt(12)
    table_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Создание таблицы
    # 8 столбцов: №, Маршрут, Отправление, Прибытие, Цена билета, Количество билетов, Общая стоимость
    table = doc.add_table(rows=8, cols=7)
    table.style = 'Light Grid Accent 1'
    
    # Заголовки столбцов
    headers = ['№', 'Маршрут', 'Отправление', 'Прибытие', 
               'Цена билета (руб.)', 'Количество проданных билетов', 'Общая стоимость (руб.)']
    
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Данные для таблицы
    routes = [
        ['1', 'Воронеж - Москва', '08:00', '14:30', '1200', '45', '=E2*F2'],
        ['2', 'Воронеж - Липецк', '09:30', '11:00', '350', '38', '=E3*F3'],
        ['3', 'Воронеж - Тамбов', '10:00', '12:30', '450', '32', '=E4*F4'],
        ['4', 'Воронеж - Курск', '11:15', '14:00', '500', '28', '=E5*F5'],
        ['5', 'Воронеж - Белгород', '13:00', '16:30', '600', '25', '=E6*F6'],
        ['6', 'Воронеж - Саратов', '14:30', '19:00', '800', '20', '=E7*F7'],
    ]
    
    # Заполнение данных
    for i, route_data in enumerate(routes, 1):
        for j, value in enumerate(route_data):
            cell = table.rows[i].cells[j]
            if j < 5:  # Номер, маршрут, время, цена
                cell.text = value
            elif j == 5:  # Количество билетов
                cell.text = value
            else:  # Общая стоимость (формула - просто умножим)
                price = int(route_data[4])
                quantity = int(route_data[5])
                cell.text = str(price * quantity)
            
            # Выравнивание
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Строка "Итого"
    itogo_row = table.rows[7]
    itogo_row.cells[0].merge(itogo_row.cells[5])
    itogo_row.cells[0].text = 'ИТОГО:'
    itogo_row.cells[0].paragraphs[0].runs[0].font.bold = True
    itogo_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Подсчет итоговой суммы
    total = sum(int(table.rows[i].cells[6].text) for i in range(1, 7))
    itogo_row.cells[6].text = str(total)
    itogo_row.cells[6].paragraphs[0].runs[0].font.bold = True
    itogo_row.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Добавление описания
    doc.add_paragraph()
    description = doc.add_paragraph()
    desc_text = """
В данной таблице представлены автобусные маршруты из города Воронеж в различные города.
Для каждого маршрута указано время отправления и прибытия, цена билета, количество 
проданных билетов и общая стоимость проданных билетов. В последней строке подведен итог -
общая сумма выручки по всем маршрутам.
"""
    description.add_run(desc_text.strip())
    description.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Сохранение документа
    filename = 'labs/Lab6_Word_Tables.docx'
    doc.save(filename)
    print(f"✓ Лабораторная работа №6 успешно создана: {filename}")
    print(f"  Создана таблица с {len(routes)} маршрутами")
    print(f"  Общая выручка: {total} руб.")
    return filename

if __name__ == "__main__":
    create_lab6_document()
