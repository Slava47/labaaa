#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Лабораторная работа №5
Microsoft Office Word. Настройки параметров документа. 
Ввод, редактирование и форматирование текста
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_number(section):
    """Добавить номера страниц"""
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def create_lab5_document():
    """Создание документа для лабораторной работы №5"""
    
    # Создание нового документа
    doc = Document()
    
    # Установка параметров страницы
    section = doc.sections[0]
    section.page_height = Cm(29.7)  # A4
    section.page_width = Cm(21)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(1)
    
    # Добавление верхнего колонтитула
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = "Ю. Левитанский. Каждый выбирает для себя."
    
    # Добавление номеров страниц
    add_page_number(section)
    
    # Заголовок стихотворения
    title = doc.add_paragraph()
    title_run = title.add_run("Каждый выбирает для себя…")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 255)  # Синий цвет
    title_run.font.underline = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(6)
    title.paragraph_format.space_after = Pt(6)
    
    # Декоративные символы (имитация)
    symbols = doc.add_paragraph()
    symbols.alignment = WD_ALIGN_PARAGRAPH.CENTER
    symbols_text = "✦ ✧ ✦ ✧ ✦ ✧ ✦ ✧ ✦ ✧ ✦"
    symbols_run = symbols.add_run(symbols_text)
    symbols_run.font.color.rgb = RGBColor(255, 0, 255)  # Фиолетовый
    
    # Текст стихотворения
    poem_text = [
        "Каждый выбирает для себя",
        "Женщину, религию, дорогу.",
        "Дьяволу служить или пророку",
        "Каждый выбирает для себя.",
        "",
        "Каждый выбирает по себе",
        "Слово для любви и для молитвы.",
        "Шпагу для дуэли, меч для битвы",
        "Каждый выбирает по себе.",
        "",
        "Каждый выбирает по себе",
        "Щит и латы. Посох и заплаты.",
        "Меру окончательной расплаты",
        "Каждый выбирает по себе",
        "",
        "Каждый выбирает для себя",
        "Выбираю тоже как умею.",
        "Ни к кому претензий не имею",
        "Каждый выбирает для себя."
    ]
    
    for line in poem_text:
        p = doc.add_paragraph()
        if line:  # Если строка не пустая
            run = p.add_run(line)
            run.font.name = 'Tahoma'
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(128, 0, 128)  # Фиолетовый
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(3)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.0
    
    # Создание нумерованного списка - Основные устройства компьютера
    doc.add_paragraph()
    list_title = doc.add_paragraph("Основные устройства компьютера:")
    list_title_run = list_title.runs[0]
    list_title_run.font.bold = True
    list_title_run.font.size = Pt(14)
    
    devices = ["системный блок;", "монитор;", "клавиатура;", "мышь."]
    for i, device in enumerate(devices, 1):
        p = doc.add_paragraph(device, style='List Number')
        p.runs[0].font.size = Pt(12)
    
    # Многоуровневый список - Любимые Web-сайты
    doc.add_paragraph()
    ml_title = doc.add_paragraph("Любимые Web-сайты")
    ml_title_run = ml_title.runs[0]
    ml_title_run.font.bold = True
    ml_title_run.font.size = Pt(14)
    
    # Уровень 1
    p1 = doc.add_paragraph("Социальная сеть:", style='List Number')
    p1.runs[0].font.size = Pt(12)
    
    # Уровень 2
    p1_1 = doc.add_paragraph("http://www.vk.com", style='List Number 2')
    p1_1.runs[0].font.size = Pt(11)
    
    p1_2 = doc.add_paragraph("http://www.facebook.com", style='List Number 2')
    p1_2.runs[0].font.size = Pt(11)
    
    p1_3 = doc.add_paragraph("http://www.instagram.com", style='List Number 2')
    p1_3.runs[0].font.size = Pt(11)
    
    # Уровень 1
    p2 = doc.add_paragraph("Музыка", style='List Number')
    p2.runs[0].font.size = Pt(12)
    
    p2_1 = doc.add_paragraph("http://www.spotify.com", style='List Number 2')
    p2_1.runs[0].font.size = Pt(11)
    
    # Уровень 1
    p3 = doc.add_paragraph("Картинки", style='List Number')
    p3.runs[0].font.size = Pt(12)
    
    p3_1 = doc.add_paragraph("http://www.pinterest.com", style='List Number 2')
    p3_1.runs[0].font.size = Pt(11)
    
    # Уровень 1
    p4 = doc.add_paragraph("Видео", style='List Number')
    p4.runs[0].font.size = Pt(12)
    
    p4_1 = doc.add_paragraph("http://www.youtube.com", style='List Number 2')
    p4_1.runs[0].font.size = Pt(11)
    
    # Сохранение документа
    filename = 'labs/Lab5_Word_Formatting.docx'
    doc.save(filename)
    print(f"✓ Лабораторная работа №5 успешно создана: {filename}")
    return filename

if __name__ == "__main__":
    create_lab5_document()
